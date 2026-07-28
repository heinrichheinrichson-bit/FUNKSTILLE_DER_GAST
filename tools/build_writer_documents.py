from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "docs" / "writer"
OUTPUT_DIR = ROOT / "exports"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Normal"].paragraph_format.line_spacing = 1.08

    for name, size, color in (
        ("Title", 24, "102B32"),
        ("Heading 1", 17, "0A7F89"),
        ("Heading 2", 13, "176B74"),
        ("Heading 3", 11, "224E55"),
    ):
        styles[name].font.name = "Aptos Display"
        styles[name].font.size = Pt(size)
        styles[name].font.color.rgb = RGBColor.from_string(color)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("FUNKSTILLE: DER GAST  ·  Writer-Dokument")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(90, 110, 114)
    return doc


def add_markdown_table(doc, lines):
    rows = []
    for line in lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", c) for c in rows[1]):
        rows.pop(1)
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = row[c_idx] if c_idx < len(row) else ""
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
                    if r_idx == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
            if r_idx == 0:
                shade_cell(cell, "176B74")
            elif r_idx % 2 == 0:
                shade_cell(cell, "EAF4F5")
    set_repeat_table_header(table.rows[0])


def add_inline_runs(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(15, 105, 115)
        else:
            paragraph.add_run(part)


def render_markdown(source, target):
    doc = setup_document()
    lines = source.read_text(encoding="utf-8").splitlines()
    i = 0
    first_title = True
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue
        if line.startswith("# "):
            if first_title:
                p = doc.add_paragraph(style="Title")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline_runs(p, line[2:])
                first_title = False
            else:
                doc.add_page_break()
                add_inline_runs(doc.add_paragraph(style="Heading 1"), line[2:])
        elif line.startswith("## "):
            p = doc.add_paragraph(style="Subtitle" if first_title is False and len(doc.paragraphs) < 4 else "Heading 2")
            if p.style.name == "Subtitle":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(p, line[3:])
        elif line.startswith("### "):
            add_inline_runs(doc.add_paragraph(style="Heading 3"), line[4:])
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.right_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(line[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(15, 105, 115)
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, line[2:])
        else:
            p = doc.add_paragraph()
            add_inline_runs(p, line)
        i += 1

    core = doc.core_properties
    core.title = source.stem.replace("_", " ")
    core.subject = "FUNKSTILLE: DER GAST – Writer-Unterlagen"
    core.author = "Vadafok Studio"
    core.keywords = "FUNKSTILLE, Story, Writer, Kontinuität, Game Design"
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(target)


def main():
    jobs = [
        (
            SOURCE_DIR / "01_REFERENZPFAD_VOLLSTAENDIGE_HANDLUNG.md",
            OUTPUT_DIR / "FUNKSTILLE_01_Vollstaendige_Handlung_Referenzpfad.docx",
        ),
        (
            SOURCE_DIR / "02_KONTINUITAET_EREIGNISSE_GEFAHREN_TODE.md",
            OUTPUT_DIR / "FUNKSTILLE_02_Kontinuitaet_Ereignisse_Gefahren_Tode.docx",
        ),
    ]
    for source, target in jobs:
        render_markdown(source, target)
        print(target)


if __name__ == "__main__":
    main()
